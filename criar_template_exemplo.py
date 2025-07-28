from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os

def create_sample_template():
    """Cria um template Word de exemplo com placeholders"""
    
    # Criar documento
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Título principal
    title = doc.add_heading('CARTA AUTOMÁTICA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adicionar espaço
    doc.add_paragraph()
    
    # Data
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run('Data: [DATA]')
    date_run.font.size = Pt(12)
    
    # Espaço
    doc.add_paragraph()
    
    # Saudação
    greeting = doc.add_paragraph('Prezado Cliente,')
    greeting.runs[0].font.bold = True
    greeting.runs[0].font.size = Pt(12)
    
    # Espaço
    doc.add_paragraph()
    
    # Corpo da carta
    body = doc.add_paragraph()
    body.add_run('Este é o seu número: ').font.size = Pt(12)
    number_run = body.add_run('[NUMERO]')
    number_run.font.bold = True
    number_run.font.size = Pt(12)
    
    body.add_run('\nE o seu ICCID: ').font.size = Pt(12)
    iccid_run = body.add_run('[ICCID]')
    iccid_run.font.bold = True
    iccid_run.font.size = Pt(12)
    
    # Espaço
    doc.add_paragraph()
    
    # Texto adicional
    additional = doc.add_paragraph('Agradecemos sua preferência e confiança em nossos serviços.')
    additional.runs[0].font.size = Pt(12)
    
    # Espaço
    doc.add_paragraph()
    
    # Assinatura
    signature = doc.add_paragraph('Atenciosamente,')
    signature.runs[0].font.size = Pt(12)
    
    company = doc.add_paragraph('Equipe de Atendimento')
    company.runs[0].font.bold = True
    company.runs[0].font.size = Pt(12)
    
    # Salvar template
    if not os.path.exists('templates_word'):
        os.makedirs('templates_word')
    
    template_path = 'templates_word/template_exemplo.docx'
    doc.save(template_path)
    
    print(f"✅ Template criado: {template_path}")
    print("📝 Placeholders disponíveis: [NUMERO], [ICCID], [DATA]")
    print("🎨 Formatação: Título centralizado, dados em negrito")
    
    return template_path

if __name__ == "__main__":
    create_sample_template() 