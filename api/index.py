from flask import Flask, render_template, request, jsonify, send_file
import pandas as pd
import os
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import io
import tempfile
import zipfile
from datetime import datetime
from docx import Document
import re
from docx2pdf import convert
import shutil
import uuid
import json
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from functools import lru_cache
import sys

# Windows-specific imports (only if available)
try:
    import pythoncom
    import win32com.client
    WINDOWS_AVAILABLE = True
except ImportError:
    WINDOWS_AVAILABLE = False
    print("⚠️ Windows dependencies not available - Word conversion will use fallback methods")

# Linux-compatible Word processing
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not available - using basic text conversion")

app = Flask(__name__)
app.config['SECRET_KEY'] = 'sua_chave_secreta_aqui'
app.config['UPLOAD_FOLDER'] = '/tmp/uploads'
app.config['TEMPLATES_FOLDER'] = '/tmp/templates_word'
app.config['TEMP_FOLDER'] = '/tmp/temp'
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max file size
app.config['MAX_WORKERS'] = 2   # Reduzido para Vercel serverless
app.config['CHUNK_SIZE'] = 2    # Chunks menores para serverless
app.config['TIMEOUT_SECONDS'] = 0  # SEM TIMEOUT - processa até terminar

# Criar pastas necessárias
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['TEMPLATES_FOLDER'], exist_ok=True)
os.makedirs(app.config['TEMP_FOLDER'], exist_ok=True)

# Dicionário para armazenar progresso
progress_tracker = {}

# Cache para templates processados
template_cache = {}

# Dicionário para armazenar tempos de início dos jobs
job_start_times = {}

# Set para controlar jobs ativos
current_jobs = set()

# Função para limpar jobs antigos
def cleanup_old_jobs():
    """Remove jobs antigos (mais de 1 hora) para evitar vazamentos de memória"""
    current_time = time.time()
    jobs_to_remove = []
    
    for job_id, job_data in progress_tracker.items():
        if current_time - job_data.get('start_time', 0) > 3600:  # 1 hora
            jobs_to_remove.append(job_id)
    
    for job_id in jobs_to_remove:
        if job_id in progress_tracker:
            del progress_tracker[job_id]
        if job_id in job_start_times:
            del job_start_times[job_id]
        current_jobs.discard(job_id)
        print(f"Job antigo removido: {job_id}")

ALLOWED_EXTENSIONS = {'xlsx', 'xls'}
ALLOWED_TEMPLATE_EXTENSIONS = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_template_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_TEMPLATE_EXTENSIONS

# Template padrão da carta (fallback)
DEFAULT_TEMPLATE = """[NOME]
[CIDADE]

[IDADE]
[PROFISSAO]
[SALARIO]

Prezado Senhor / Senhora,

Espero que esta carta o encontre bem. Eu queria escrever para você para [DEPARTAMENTO]. Eu sinto que é importante compartilhar meus pensamentos sobre este assunto.

Com os melhores cumprimentos,
[NOME]"""

@app.route('/api/health')
def health_check():
    """Endpoint de verificação de saúde da aplicação"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'active_jobs': len(current_jobs),
        'workers': app.config['MAX_WORKERS'],
        'version': '1.0.0'
    })

@app.route('/')
def index():
    # Listar templates disponíveis
    templates = []
    if os.path.exists(app.config['TEMPLATES_FOLDER']):
        for file in os.listdir(app.config['TEMPLATES_FOLDER']):
            if file.endswith('.docx'):
                templates.append(file)
    
    return render_template('index.html', templates=templates)

@app.route('/api/upload-template', methods=['POST'])
def upload_template():
    """Upload de template Word"""
    if 'template' not in request.files:
        return jsonify({'error': 'Nenhum arquivo enviado'}), 400
    
    file = request.files['template']
    if file.filename == '':
        return jsonify({'error': 'Nenhum arquivo selecionado'}), 400
    
    if file and allowed_template_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['TEMPLATES_FOLDER'], filename)
        file.save(filepath)
        return jsonify({'message': 'Template enviado com sucesso', 'filename': filename})
    
    return jsonify({'error': 'Tipo de arquivo não permitido'}), 400

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Upload de arquivo Excel"""
    if 'file' not in request.files:
        return jsonify({'error': 'Nenhum arquivo enviado'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nenhum arquivo selecionado'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Ler dados do Excel
        try:
            df = pd.read_excel(filepath)
            data = df.to_dict('records')
            return jsonify({
                'message': 'Arquivo enviado com sucesso',
                'filename': filename,
                'records': len(data),
                'columns': list(df.columns)
            })
        except Exception as e:
            return jsonify({'error': f'Erro ao ler arquivo: {str(e)}'}), 400
    
    return jsonify({'error': 'Tipo de arquivo não permitido'}), 400

@app.route('/api/generate-pdf', methods=['POST'])
def generate_pdf():
    """Gerar PDFs a partir de dados Excel e template Word"""
    try:
        data = request.get_json()
        excel_file = data.get('excel_file')
        template_name = data.get('template_name')
        use_word_template = data.get('use_word_template', True)
        
        if not excel_file or not template_name:
            return jsonify({'error': 'Arquivo Excel e template são obrigatórios'}), 400
        
        # Ler dados do Excel
        excel_path = os.path.join(app.config['UPLOAD_FOLDER'], excel_file)
        if not os.path.exists(excel_path):
            return jsonify({'error': 'Arquivo Excel não encontrado'}), 404
        
        df = pd.read_excel(excel_path)
        data_list = df.to_dict('records')
        
        if len(data_list) == 0:
            return jsonify({'error': 'Nenhum dado encontrado no arquivo Excel'}), 400
        
        # Gerar job ID
        job_id = str(uuid.uuid4())
        current_jobs.add(job_id)
        
        # Inicializar progresso
        progress_tracker[job_id] = {
            'total': len(data_list),
            'completed': 0,
            'current': 0,
            'status': 'starting',
            'start_time': time.time(),
            'rate': 0,
            'estimated_time': 0
        }
        job_start_times[job_id] = time.time()
        
        # Iniciar processamento em background
        thread = threading.Thread(
            target=generate_multiple_pdfs_parallel,
            args=(data_list, template_name, use_word_template, job_id),
            daemon=True
        )
        thread.start()
        
        return jsonify({
            'message': 'Geração iniciada',
            'job_id': job_id,
            'total_records': len(data_list)
        })
        
    except Exception as e:
        return jsonify({'error': f'Erro interno: {str(e)}'}), 500

@app.route('/api/progress/<job_id>')
def get_progress(job_id):
    """Obter progresso da geração"""
    if job_id not in progress_tracker:
        return jsonify({'error': 'Job não encontrado'}), 404
    
    progress = progress_tracker[job_id]
    elapsed_time = time.time() - job_start_times.get(job_id, time.time())
    
    return jsonify({
        'job_id': job_id,
        'total': progress['total'],
        'completed': progress['completed'],
        'current': progress['current'],
        'status': progress['status'],
        'elapsed_time': elapsed_time,
        'rate': progress['rate'],
        'estimated_time': progress['estimated_time']
    })

@app.route('/api/download/<job_id>')
def download_result(job_id):
    """Download do resultado"""
    if job_id not in progress_tracker:
        return jsonify({'error': 'Job não encontrado'}), 404
    
    progress = progress_tracker[job_id]
    if progress['status'] != 'completed':
        return jsonify({'error': 'Job ainda não foi concluído'}), 400
    
    zip_path = progress.get('zip_path')
    if not zip_path or not os.path.exists(zip_path):
        return jsonify({'error': 'Arquivo ZIP não encontrado'}), 404
    
    return send_file(zip_path, as_attachment=True, download_name=f'result_{job_id}.zip')

def generate_multiple_pdfs_parallel(data_list, template_name, use_word_template, job_id):
    """Gerar múltiplos PDFs em paralelo"""
    try:
        print(f"🚀 Iniciando geração otimizada de PDFs usando template '{template_name}'...")
        
        # Atualizar status
        progress_tracker[job_id]['status'] = 'processing'
        
        # Dividir dados em chunks
        chunk_size = app.config['CHUNK_SIZE']
        chunks = [data_list[i:i + chunk_size] for i in range(0, len(data_list), chunk_size)]
        
        pdf_files = []
        total_chunks = len(chunks)
        
        # Processar chunks em paralelo
        with ThreadPoolExecutor(max_workers=app.config['MAX_WORKERS']) as executor:
            futures = []
            
            for chunk_id, chunk in enumerate(chunks):
                future = executor.submit(
                    process_chunk_optimized,
                    chunk, template_name, use_word_template, job_id, chunk_id
                )
                futures.append(future)
            
            # Coletar resultados
            for future in as_completed(futures):
                try:
                    chunk_pdfs = future.result()
                    pdf_files.extend(chunk_pdfs)
                    
                    # Atualizar progresso
                    progress_tracker[job_id]['completed'] += len(chunk_pdfs)
                    progress_tracker[job_id]['current'] = progress_tracker[job_id]['completed']
                    
                    # Calcular taxa e tempo estimado
                    elapsed = time.time() - job_start_times[job_id]
                    if elapsed > 0:
                        rate = progress_tracker[job_id]['completed'] / elapsed
                        remaining = (progress_tracker[job_id]['total'] - progress_tracker[job_id]['completed']) / rate if rate > 0 else 0
                        
                        progress_tracker[job_id]['rate'] = rate
                        progress_tracker[job_id]['estimated_time'] = remaining
                    
                    print(f"⚡ SUPER Progresso: {progress_tracker[job_id]['completed']}/{progress_tracker[job_id]['total']} ({rate:.2f}/s)")
                    
                except Exception as e:
                    print(f"❌ Erro no chunk: {e}")
                    continue
        
        # Criar ZIP final
        if pdf_files:
            zip_path = os.path.join(app.config['TEMP_FOLDER'], f'result_{job_id}.zip')
            create_final_zip(pdf_files, zip_path)
            
            # Atualizar status final
            progress_tracker[job_id]['status'] = 'completed'
            progress_tracker[job_id]['zip_path'] = zip_path
            
            total_time = time.time() - job_start_times[job_id]
            print(f"🎉 SUPER ULTRA: {len(pdf_files)} PDFs em {total_time:.1f}s")
            print(f"🚀 Taxa SUPER: {len(pdf_files)/total_time:.2f} PDFs/segundo")
            print(f"📁 ZIP salvo em: {zip_path}")
        
        # Limpar jobs antigos
        cleanup_old_jobs()
        
    except Exception as e:
        print(f"❌ Erro na geração: {e}")
        progress_tracker[job_id]['status'] = 'error'
        progress_tracker[job_id]['error'] = str(e)
    finally:
        current_jobs.discard(job_id)

def process_chunk_optimized(chunk, template_name, use_word_template, job_id, chunk_id):
    """Processar chunk de dados otimizado"""
    pdf_files = []
    
    for i, row_data in enumerate(chunk):
        try:
            # Gerar nome do arquivo
            numero = row_data.get('NUMERO', f"{i+1:03d}")
            filename = f"Carta_{numero}.pdf"
            
            # Gerar PDF
            if use_word_template:
                pdf_content = generate_word_pdf_ultra_optimized(row_data, template_name)
            else:
                pdf_content = generate_simple_pdf(row_data, DEFAULT_TEMPLATE)
            
            # Salvar PDF
            pdf_path = os.path.join(app.config['TEMP_FOLDER'], filename)
            with open(pdf_path, 'wb') as f:
                f.write(pdf_content)
            
            pdf_files.append(pdf_path)
            print(f"      ✅ PDF gerado: {filename}")
            
        except Exception as e:
            print(f"❌ Erro ao gerar PDF {i+1}: {e}")
            continue
    
    print(f"   ✅ Chunk {chunk_id} concluído: {len(pdf_files)} PDFs gerados")
    return pdf_files

def create_final_zip(pdf_files, zip_path):
    """Criar ZIP final com todos os PDFs"""
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for pdf_file in pdf_files:
            if os.path.exists(pdf_file):
                zipf.write(pdf_file, os.path.basename(pdf_file))
    
    print(f"📦 Criando ZIP final...")

def cleanup_temp_files(pdf_files):
    """Limpar arquivos temporários"""
    for pdf_file in pdf_files:
        try:
            if os.path.exists(pdf_file):
                os.remove(pdf_file)
        except:
            pass

@lru_cache(maxsize=10)
def prepare_template_cache(template_name):
    """Preparar cache do template"""
    template_path = os.path.join(app.config['TEMPLATES_FOLDER'], template_name)
    if os.path.exists(template_path):
        return template_path
    return None

def generate_simple_pdf(row_data, template_text):
    """Gera PDF simples com ReportLab"""
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
    story = []
    
    # Substituir placeholders
    content = template_text
    for key, value in row_data.items():
        placeholder = f'[{key.upper()}]'
        content = content.replace(placeholder, str(value) if value is not None else '')
    
    # Adicionar conteúdo ao PDF
    styles = getSampleStyleSheet()
    story.append(Paragraph(content, styles['Normal']))
    
    doc.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer.getvalue()

def generate_word_pdf_ultra_optimized(row_data, template_name):
    """Gera PDF a partir de template Word ultra otimizado"""
    try:
        # Obter template do cache
        template_path = prepare_template_cache(template_name)
        if not template_path:
            raise Exception(f"Template não encontrado: {template_name}")
        
        # Criar arquivos temporários
        temp_docx = os.path.join(app.config['TEMP_FOLDER'], f"temp_{uuid.uuid4()}.docx")
        temp_pdf = os.path.join(app.config['TEMP_FOLDER'], f"temp_{uuid.uuid4()}.pdf")
        
        try:
            # Copiar template
            shutil.copy2(template_path, temp_docx)
            
            # Substituir placeholders
            doc = Document(temp_docx)
            
            def replace_placeholders_robust():
                # Mapeamento de placeholders longos específicos
                placeholder_mapping = {
                    '[NUMERO]': 'NUMERO',
                    '[ICCID]': 'ICCID',
                    '[NOME]': 'NOME',
                    '[CIDADE]': 'CIDADE',
                    '[IDADE]': 'IDADE',
                    '[PROFISSAO]': 'PROFISSAO',
                    '[SALARIO]': 'SALARIO',
                    '[DEPARTAMENTO]': 'DEPARTAMENTO'
                }
                
                # Substituir placeholders específicos primeiro
                for placeholder, key in placeholder_mapping.items():
                    if key in row_data:
                        value = row_data[key]
                        if value is not None:
                            print(f"   🔄 Substituído placeholder específico: {placeholder} → {value}")
                
                # Substituir placeholders dinâmicos (colunas do Excel)
                for key, value in row_data.items():
                    if value is not None:
                        placeholder = f'[{key.upper()}]'
                        # Substituir em todo o documento
                        for paragraph in doc.paragraphs:
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(value))
                        
                        # Substituir em tabelas
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    if placeholder in cell.text:
                                        cell.text = cell.text.replace(placeholder, str(value))
            
            def preserve_document_formatting():
                # Preservar estilos de parágrafo
                for paragraph in doc.paragraphs:
                    # Manter alinhamento original
                    original_alignment = paragraph.alignment
                    
                    # Preservar formatação de runs
                    for run in paragraph.runs:
                        # Manter negrito, itálico, sublinhado
                        original_bold = run.bold
                        original_italic = run.italic
                        original_underline = run.underline
                        
                        # Aplicar formatação preservada
                        if original_bold is not None:
                            run.bold = original_bold
                        if original_italic is not None:
                            run.italic = original_italic
                        if original_underline is not None:
                            run.underline = original_underline
                    
                    # Restaurar alinhamento
                    if original_alignment is not None:
                        paragraph.alignment = original_alignment
                
                # Preservar formatação de tabelas
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                # Preservar alinhamento das células
                                original_alignment = paragraph.alignment
                                if original_alignment is not None:
                                    paragraph.alignment = original_alignment
                
                print(f"   🔄 Substituído placeholder preservando formatação completa")
            
            # Executar substituições
            replace_placeholders_robust()
            preserve_document_formatting()
            
            # Salvar documento modificado
            doc.save(temp_docx)
            
            # Converter para PDF
            print(f"🔄 Iniciando conversão Word→PDF: {os.path.basename(temp_docx)}")
            
            # Tentar diferentes métodos de conversão
            try:
                # Método 1: docx2pdf
                print(f"   📄 Tentando docx2pdf...")
                convert(temp_docx, temp_pdf)
                print(f"   ✅ Conversão bem-sucedida com docx2pdf")
            except Exception as e:
                print(f"   ❌ docx2pdf falhou: {e}")
                
                # Método 2: Fallback
                print(f"   🖥️ Tentando COM direto...")
                convert_word_to_pdf_com_preserve_formatting(temp_docx, temp_pdf)
            
            # Verificar se PDF foi criado
            if os.path.exists(temp_pdf):
                with open(temp_pdf, 'rb') as f:
                    pdf_content = f.read()
                return pdf_content
            else:
                raise Exception("PDF não foi criado")
                
        finally:
            # Limpar arquivos temporários
            try:
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)
                if os.path.exists(temp_pdf):
                    os.remove(temp_pdf)
            except:
                pass
                
    except Exception as e:
        print(f"❌ Erro na geração Word PDF: {e}")
        # Fallback para PDF simples
        return generate_simple_pdf(row_data, DEFAULT_TEMPLATE)

def convert_word_to_pdf_com_preserve_formatting(docx_path, pdf_path):
    """Converte Word para PDF usando COM direto preservando formatação exata"""
    
    # Fallback se Windows não estiver disponível
    if not WINDOWS_AVAILABLE:
        print(f"   ⚠️ Windows não disponível, usando fallback")
        return convert_word_to_pdf_fallback(docx_path, pdf_path)
    
    word = None
    doc = None
    try:
        # Inicializar COM para esta thread
        pythoncom.CoInitialize()
        
        # Criar instância do Word com DispatchEx para nova instância
        word = win32com.client.DispatchEx("Word.Application")
        
        # Configurar Word para preservar formatação
        word.Visible = False
        word.DisplayAlerts = False
        
        # Verificar se arquivo existe
        if not os.path.exists(docx_path):
            raise Exception(f"Arquivo não encontrado: {docx_path}")
        
        # Abrir documento com caminho absoluto
        abs_docx_path = os.path.abspath(docx_path)
        abs_pdf_path = os.path.abspath(pdf_path)
        
        print(f"   📄 Abrindo documento: {os.path.basename(docx_path)}")
        doc = word.Documents.Open(abs_docx_path)
        
        # Configurar opções de PDF para preservar formatação
        print(f"   🔄 Convertendo para PDF...")
        
        # Múltiplas tentativas com diferentes parâmetros (versão robusta)
        methods = [
            lambda: doc.SaveAs(abs_pdf_path, FileFormat=17, OptimizeFor=0),
            lambda: doc.SaveAs(abs_pdf_path, FileFormat=17),
            lambda: doc.SaveAs(abs_pdf_path, FileFormat=6),
            lambda: doc.SaveAs(abs_pdf_path)
        ]
        
        success = False
        for i, method in enumerate(methods):
            try:
                method()
                success = True
                print(f"   ✅ Conversão bem-sucedida com método {i+1}")
                break
            except Exception as e:
                if i == len(methods) - 1:  # Última tentativa
                    raise e
                print(f"   ⚠️ Método {i+1} falhou, tentando próximo...")
                continue
        
        # Fechar documento
        doc.Close(False)  # False = não salvar alterações
        
        # Verificar se PDF foi criado
        if os.path.exists(pdf_path):
            file_size = os.path.getsize(pdf_path)
            if file_size > 1000:  # PDF deve ter pelo menos 1KB
                print(f"   ✅ PDF criado com sucesso: {file_size} bytes")
                with open(pdf_path, 'rb') as f:
                    pdf_content = f.read()
                return pdf_content
            else:
                raise Exception(f"PDF criado mas muito pequeno: {file_size} bytes")
        else:
            raise Exception("PDF não foi criado")
            
    except Exception as e:
        print(f"   ❌ Erro na conversão COM: {e}")
        raise e
        
    finally:
        # Limpar recursos de forma robusta
        if doc is not None:
            try:
                doc.Close(False)
            except:
                pass
        if word is not None:
            try:
                word.Quit()
            except:
                pass
        try:
            pythoncom.CoUninitialize()
        except:
            pass

def convert_word_to_pdf_fallback(docx_path, pdf_path):
    """Fallback para conversão Word para PDF quando Windows não está disponível"""
    try:
        print(f"   📄 Usando fallback para conversão")
        
        # Ler o documento Word
        doc = Document(docx_path)
        
        # Gerar PDF com formatação melhorada
        pdf_buffer = io.BytesIO()
        doc_pdf = SimpleDocTemplate(pdf_buffer, pagesize=A4, 
                                   topMargin=0.5*inch, bottomMargin=0.5*inch,
                                   leftMargin=0.5*inch, rightMargin=0.5*inch)
        story = []
        
        # Estilos melhorados
        styles = getSampleStyleSheet()
        
        # Estilo para logo DIGI
        logo_style = ParagraphStyle(
            'Logo',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=25,
            alignment=1,  # Centralizado
            textColor=colors.HexColor('#0915FF'),
            fontName='Helvetica-Bold'
        )
        
        # Estilo para subtítulos
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            alignment=0,  # Esquerda
            textColor=colors.black,
            fontName='Helvetica-Bold'
        )
        
        # Estilo para texto normal
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=8,
            leading=14,
            alignment=0,  # Esquerda
            textColor=colors.black
        )
        
        # Estilo para dados da tabela
        data_style = ParagraphStyle(
            'Data',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leading=12,
            alignment=0,  # Esquerda
            textColor=colors.black,
            fontName='Helvetica'
        )
        
        # Processar parágrafos preservando formatação
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Processar cada run para preservar formatação
                paragraph_content = []
                
                for run in paragraph.runs:
                    text = run.text
                    is_bold = run.bold
                    is_underline = run.underline
                    
                    # Aplicar formatação baseada no conteúdo e estilo
                    if 'digi' in text.lower() and len(text.strip()) <= 10:
                        # Logo DIGI
                        paragraph_content.append(f'<font color="#0915FF" size="18"><b>{text}</b></font>')
                    elif is_bold and any(keyword in text.lower() for keyword in ['bem-vindo', 'número', 'iccid', 'dúvida', 'vantagens']):
                        # Texto importante em negrito
                        paragraph_content.append(f'<b>{text}</b>')
                    elif is_underline:
                        # Texto sublinhado (cabeçalhos de tabela)
                        paragraph_content.append(f'<u>{text}</u>')
                    elif is_bold:
                        # Texto em negrito
                        paragraph_content.append(f'<b>{text}</b>')
                    else:
                        # Texto normal
                        paragraph_content.append(text)
                
                # Juntar conteúdo do parágrafo
                full_text = ''.join(paragraph_content)
                
                # Determinar estilo baseado no conteúdo
                if 'digi' in full_text.lower() and len(full_text.strip()) <= 10:
                    # Logo centralizado
                    story.append(Paragraph(full_text, logo_style))
                    story.append(Spacer(1, 20))
                elif any(keyword in full_text.lower() for keyword in ['bem-vindo', 'número', 'iccid']):
                    # Subtítulos importantes
                    story.append(Paragraph(full_text, subtitle_style))
                    story.append(Spacer(1, 12))
                elif any(keyword in full_text.lower() for keyword in ['contatar', 'contactar', 'dúvida', 'ajudar']):
                    # Seção de contato
                    story.append(Paragraph(full_text, normal_style))
                    story.append(Spacer(1, 8))
                else:
                    # Texto normal
                    story.append(Paragraph(full_text, normal_style))
                    story.append(Spacer(1, 8))
        
        # Processar tabelas com formatação melhorada
        for table in doc.tables:
            table_data = []
            has_header = False
            
            for i, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    # Processar formatação das células
                    cell_content = []
                    for run in cell.paragraphs[0].runs:
                        if run.underline:
                            cell_content.append(f'<u>{run.text}</u>')
                        elif run.bold:
                            cell_content.append(f'<b>{run.text}</b>')
                        else:
                            cell_content.append(run.text)
                    
                    cell_text = ''.join(cell_content).strip()
                    row_data.append(cell_text)
                    
                    # Detectar se é cabeçalho (primeira linha com texto sublinhado)
                    if i == 0 and any(run.underline for run in cell.paragraphs[0].runs):
                        has_header = True
                
                table_data.append(row_data)
            
            if table_data:
                # Criar tabela no PDF
                pdf_table = Table(table_data)
                
                if has_header:
                    # Tabela com cabeçalho formatado
                    pdf_table.setStyle(TableStyle([
                        # Cabeçalho da tabela
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0915FF')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                        # Corpo da tabela
                        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                        ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 11),
                        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#0915FF')),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F8FF')])
                    ]))
                else:
                    # Tabela simples
                    pdf_table.setStyle(TableStyle([
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 0), (-1, -1), 11),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('BACKGROUND', (0, 0), (-1, -1), colors.white)
                    ]))
                
                story.append(pdf_table)
                story.append(Spacer(1, 15))
        
        doc_pdf.build(story)
        pdf_buffer.seek(0)
        
        # Salvar PDF
        with open(pdf_path, 'wb') as f:
            f.write(pdf_buffer.getvalue())
        
        file_size = os.path.getsize(pdf_path)
        print(f"   ✅ PDF fallback criado: {file_size} bytes")
        
        return pdf_buffer.getvalue()
        
    except Exception as e:
        print(f"   ❌ Erro no fallback: {e}")
        raise e

# Para Vercel serverless
app.debug = False 